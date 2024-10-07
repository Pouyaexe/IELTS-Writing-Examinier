import markdown
from docx import Document
from bs4 import BeautifulSoup
import io

def convert_markdown_to_docx(markdown_text):
    """Convert a markdown text to a DOCX file and return it as a binary object."""
    html_string = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html_string, "html.parser")
    doc = Document()

    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "p":
            doc.add_paragraph(element.text)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer
